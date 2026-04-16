from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "finalreport.md"
OUTPUT_DOCX = ROOT / "finalreport.docx"

IMAGE_DIRS = [
    ROOT / "docs" / "RM&EERD",
    ROOT / "docs" / "platform-pictures",
    ROOT / "docs" / "supabase-pictures",
]

PLACEHOLDER_RE = re.compile(r"^\[ADD (.+?) PICTURE HERE\]$")
ORDERED_ITEM_RE = re.compile(r"^(\d+)\.\s+(.*)$")
BULLET_ITEM_RE = re.compile(r"^-\s+(.*)$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
URL_RE = re.compile(r"https?://\S+")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|https?://\S+)")

SQL_KEYWORDS = {
    "add",
    "after",
    "all",
    "alter",
    "and",
    "any",
    "as",
    "asc",
    "audit_logs",
    "by",
    "bigint",
    "bigserial",
    "build_object",
    "cascade",
    "case",
    "coalesce",
    "column",
    "count",
    "create",
    "default",
    "delete",
    "desc",
    "distinct",
    "do",
    "exists",
    "excluded",
    "false",
    "filter",
    "for",
    "foreign",
    "from",
    "function",
    "gen_random_uuid",
    "group",
    "if",
    "in",
    "index",
    "insert",
    "into",
    "is",
    "join",
    "jsonb",
    "jsonb_build_object",
    "key",
    "language",
    "left",
    "limit",
    "not",
    "now",
    "null",
    "on",
    "or",
    "order",
    "primary",
    "references",
    "replace",
    "restrict",
    "returning",
    "returns",
    "round",
    "select",
    "set",
    "sum",
    "table",
    "text",
    "then",
    "timestamptz",
    "trigger",
    "true",
    "union",
    "unique",
    "update",
    "values",
    "view",
    "void",
    "where",
}


def set_paragraph_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = p_pr.find(qn("w:pBdr"))
    if p_borders is None:
        p_borders = OxmlElement("w:pBdr")
        p_pr.append(p_borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    p_borders.append(bottom)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, name: str, size_pt: float, color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    title = document.styles["Title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title.font.size = Pt(20)
    title.font.bold = True

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Times New Roman"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    subtitle.font.size = Pt(15)
    subtitle.font.bold = True

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0x20, 0x3A, 0x5B)

    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading2.font.size = Pt(13)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0x20, 0x3A, 0x5B)


def resolve_image_path(name: str) -> Path:
    if name.endswith(".drawio"):
        candidate = ROOT / "docs" / "RM&EERD" / f"{name}.png"
        if candidate.exists():
            return candidate
    for image_dir in IMAGE_DIRS:
        candidate = image_dir / name
        if candidate.exists():
            return candidate
    raise FileNotFoundError(f"Could not find image for placeholder: {name}")


def next_non_empty(lines: list[str], start: int) -> str:
    for idx in range(start, len(lines)):
        if lines[idx].strip():
            return lines[idx].strip()
    return ""


def add_hyperlink_like_run(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.underline = True
    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)


def add_inline_runs(paragraph, text: str) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Consolas", 10)
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
        else:
            add_hyperlink_like_run(paragraph, token)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def split_markdown_row(line: str) -> list[str]:
    parts = [part.strip() for part in line.strip().strip("|").split("|")]
    return parts


def add_table(document: Document, table_lines: list[str]) -> None:
    rows = [split_markdown_row(line) for line in table_lines]
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[2:]
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, cell_text in enumerate(header):
        cell = header_cells[idx]
        set_cell_shading(cell, "D9E2F3")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(paragraph, cell_text)
        for run in paragraph.runs:
            run.bold = True
    for row_data in body:
        row = table.add_row().cells
        for idx, cell_text in enumerate(row_data):
            paragraph = row[idx].paragraphs[0]
            add_inline_runs(paragraph, cell_text)


def add_image(document: Document, image_path: Path) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    width = 6.5 if image_path.parent.name in {"RM&EERD", "supabase-pictures"} else 6.15
    run.add_picture(str(image_path), width=Inches(width))


def add_horizontal_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    set_paragraph_border(paragraph)


def add_heading(document: Document, level: int, text: str, heading_index: int, in_front_matter: bool) -> None:
    if level == 1 and heading_index == 0:
        paragraph = document.add_paragraph(style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(text)
        return
    if level == 2 and in_front_matter and text == "Final Report":
        paragraph = document.add_paragraph(style="Subtitle")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(text)
        return
    style = "Heading 1" if level == 2 else "Heading 2"
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)


def tokenize_sql(line: str) -> list[tuple[str, str]]:
    tokens: list[tuple[str, str]] = []
    idx = 0
    while idx < len(line):
        if line[idx:].startswith("--"):
            tokens.append(("comment", line[idx:]))
            break
        if line[idx].isspace():
            start = idx
            while idx < len(line) and line[idx].isspace():
                idx += 1
            tokens.append(("plain", line[start:idx]))
            continue
        if line[idx] == "'":
            start = idx
            idx += 1
            while idx < len(line):
                if line[idx] == "'" and (idx + 1 >= len(line) or line[idx + 1] != "'"):
                    idx += 1
                    break
                idx += 1
            tokens.append(("string", line[start:idx]))
            continue
        if line[idx].isdigit():
            start = idx
            while idx < len(line) and (line[idx].isdigit() or line[idx] == "."):
                idx += 1
            tokens.append(("number", line[start:idx]))
            continue
        if line[idx].isalpha() or line[idx] == "_":
            start = idx
            while idx < len(line) and (line[idx].isalnum() or line[idx] in {"_", "."}):
                idx += 1
            word = line[start:idx]
            token_type = "keyword" if word.lower() in SQL_KEYWORDS else "plain"
            tokens.append((token_type, word))
            continue
        tokens.append(("plain", line[idx]))
        idx += 1
    return tokens


def add_code_block(document: Document, code_lines: list[str], language: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F5F5")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    for line_index, line in enumerate(code_lines):
        if language.lower() == "sql":
            for token_type, token in tokenize_sql(line):
                run = paragraph.add_run(token)
                set_run_font(run, "Consolas", 9.5)
                if token_type == "keyword":
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0xAA)
                elif token_type == "string":
                    run.font.color.rgb = RGBColor(0xA3, 0x15, 0x15)
                elif token_type == "comment":
                    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                elif token_type == "number":
                    run.font.color.rgb = RGBColor(0x09, 0x84, 0xA3)
        else:
            run = paragraph.add_run(line)
            set_run_font(run, "Consolas", 9.5)
        if line_index < len(code_lines) - 1:
            paragraph.add_run().add_break()


def build_docx(source_text: str) -> tuple[Document, int]:
    document = Document()
    configure_styles(document)

    lines = source_text.splitlines()
    idx = 0
    heading_index = 0
    inserted_images = 0
    in_front_matter = True
    pending_new_section = False

    def maybe_insert_section_break() -> None:
        nonlocal pending_new_section
        if pending_new_section:
            document.add_section(WD_SECTION.NEW_PAGE)
            pending_new_section = False

    while idx < len(lines):
        raw_line = lines[idx]
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            maybe_insert_section_break()
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            add_heading(document, level, text, heading_index, in_front_matter)
            heading_index += 1
            if text == "Abstract":
                in_front_matter = False
            idx += 1
            continue

        if stripped == "---":
            upcoming = next_non_empty(lines, idx + 1)
            if upcoming == "## Abstract" or re.match(r"^##\s+\d+\.", upcoming):
                pending_new_section = True
            else:
                add_horizontal_rule(document)
            idx += 1
            continue

        placeholder_match = PLACEHOLDER_RE.match(stripped)
        if placeholder_match:
            maybe_insert_section_break()
            image_path = resolve_image_path(placeholder_match.group(1))
            add_image(document, image_path)
            inserted_images += 1
            idx += 1
            continue

        if stripped.startswith("```"):
            maybe_insert_section_break()
            language = stripped[3:].strip()
            idx += 1
            code_lines: list[str] = []
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx].rstrip("\n"))
                idx += 1
            add_code_block(document, code_lines, language)
            if idx < len(lines):
                idx += 1
            continue

        if stripped.startswith("|") and idx + 1 < len(lines) and lines[idx + 1].strip().startswith("|"):
            maybe_insert_section_break()
            table_lines = [stripped]
            idx += 1
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            add_table(document, table_lines)
            continue

        ordered_match = ORDERED_ITEM_RE.match(stripped)
        if ordered_match:
            maybe_insert_section_break()
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, ordered_match.group(2))
            idx += 1
            continue

        bullet_match = BULLET_ITEM_RE.match(stripped)
        if bullet_match:
            maybe_insert_section_break()
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, bullet_match.group(1))
            idx += 1
            continue

        maybe_insert_section_break()
        paragraph_lines = [stripped]
        idx += 1
        while idx < len(lines):
            lookahead = lines[idx].strip()
            if (
                not lookahead
                or lookahead == "---"
                or HEADING_RE.match(lookahead)
                or PLACEHOLDER_RE.match(lookahead)
                or lookahead.startswith("```")
                or ORDERED_ITEM_RE.match(lookahead)
                or BULLET_ITEM_RE.match(lookahead)
                or (lookahead.startswith("|") and idx + 1 < len(lines) and lines[idx + 1].strip().startswith("|"))
            ):
                break
            paragraph_lines.append(lookahead)
            idx += 1

        paragraph = document.add_paragraph()
        if in_front_matter:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(paragraph, " ".join(paragraph_lines))

    return document, inserted_images


def main() -> None:
    source_text = SOURCE_MD.read_text(encoding="utf-8")
    document, inserted_images = build_docx(source_text)
    document.save(str(OUTPUT_DOCX))
    print(f"Wrote {OUTPUT_DOCX}")
    print(f"Inserted {inserted_images} images")


if __name__ == "__main__":
    main()
